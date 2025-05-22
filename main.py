import marimo

__generated_with = "0.13.11"
app = marimo.App(width="medium")


@app.cell
def _():
    import marimo as mo
    return


@app.cell
def _():
    import tomllib
    import numpy as np
    import matplotlib.pyplot as plt

    plt.rcParams["font.family"] = "Times New Roman"

    with open("./data.toml", "rb") as f:
        data = tomllib.load(f)
    return data, np, plt


@app.cell
def _(data, np, plt):
    years = range(1938, 1944)
    articulos = [len([article for article in data["articles"] if article["date"].year == year ]) for year in years]
    portadas = [len([figure for figure in data["figures"] if figure["date"].year == year and "portada" in figure ]) for year in years]
    ilustraciones = [len([figure for figure in data["figures"] if figure["date"].year == year and "portada" not in figure ]) for year in years]

    total = sum(articulos) + sum(portadas) + sum(ilustraciones)

    # Set width of bars
    bar_width = 0.25

    # Set positions of bars on x-axis
    r1 = np.arange(len(years))
    r2 = [x + bar_width for x in r1]
    r3 = [x + bar_width for x in r2]

    # Create the grouped bar chart
    plt.figure(figsize=(10, 6), dpi=300)
    bar1 = plt.bar(r1, articulos, width=bar_width, label='Artículos')
    bar2 = plt.bar(r2, portadas, width=bar_width, label='Portadas')
    bar3 = plt.bar(r3, ilustraciones, width=bar_width, label='Ilustraciones')

    def add_labels(bars):
        for bar in bars:
            height = bar.get_height()
            if height == 0:
                continue
            plt.text(bar.get_x() + bar.get_width()/2., height/2,
                    f'{height}',
                    ha='center', va='center', color='white', fontweight='bold')

    # Add value labels to each bar
    add_labels(bar1)
    add_labels(bar2)
    add_labels(bar3)

    # Add labels and title
    plt.title(f'(Total: {total})')
    plt.xlabel('Año')
    plt.xticks([r + bar_width for r in range(len(years))], years)
    plt.legend()

    # Display the chart
    plt.tight_layout()

    plt.savefig('types.png', dpi=300)
    plt.savefig('types.svg', dpi=300)
    return


@app.cell(hide_code=True)
def _(data, plt):
    from collections import Counter

    # Function to extract categories from the data
    def extract_categories(data):
        categories = []

        # Check if 'articles' exists in the data
        if 'articles' in data:
            # Iterate through each article
            for article in data['articles']:
                # Check if the article has a 'categoria' field
                if 'categoria' in article:
                    categories.append(article['categoria'])

        return categories



    # Extract categories from the data
    raw_categories = extract_categories(data)

    categories = list(set(raw_categories))

    # Count the occurrences of each category
    category_sizes = [len([c for c in raw_categories if c == category]) for category in categories]

    category_sizes, categories = zip(*sorted(zip(category_sizes, categories), reverse=True))

    # Create the pie chart tp
    plt.figure(figsize=(10, 7), dpi=300)

    colors = plt.cm.Set3(range(len(category_sizes)))
    explode = [0.05] * len(category_sizes)

    plt.pie(category_sizes, labels=categories, autopct='%1.1f%%', startangle=0, colors=colors, explode=explode)
    plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle

    plt.tight_layout()
    plt.savefig('categories.png', dpi=300)
    plt.savefig('categories.svg', dpi=300)

    return Counter, categories, colors


@app.cell(hide_code=True)
def _(Counter, categories, colors, data, plt):
    from collections import defaultdict

    # Assuming 'data' is your parsed TOML data
    # For demonstration, I'll show how to process it

    def create_archaeology_chart(data):
        # Extract all items that have a 'categoria' field (only articles have this)
        articles = data.get('articles', [])
    
        # Separate items by archaeology status
        archaeology_true = []
        archaeology_false = []
    
        for article in articles:
            if article.get('arqueologia', False):
                archaeology_true.append(article.get('categoria', 'Unknown'))
            else:
                archaeology_false.append(article.get('categoria', 'Unknown'))
    
        # Count categories for each group
        true_counts = Counter(archaeology_true)
        false_counts = Counter(archaeology_false)
    
        # Get all unique categories
        all_categories = categories
        print(all_categories)
        # all_categories = set(true_counts.keys()) | set(false_counts.keys())
        # all_categories = sorted(list(all_categories))  # Sort for consistent ordering
    
        # Prepare data for stacked bar chart
        true_values = [true_counts.get(cat, 0) for cat in all_categories]
        false_values = [false_counts.get(cat, 0) for cat in all_categories]
    
        # Create the plot
        fig, ax = plt.subplots(figsize=(12, 8), dpi=300)
    
        # Set up the bars
        bar_width = 0.6
        x_positions = [0, 1]
        bar_labels = ['Artículos arqueológicos', 'Artículos no arqueológicos']
    
        # Create distinct color palette for many categories
        # if len(all_categories) <= 10:
        #     colors = plt.cm.tab10(np.linspace(0, 1, len(all_categories)))
        # elif len(all_categories) <= 20:
            # Combine tab10 and tab20 for better distinction
        # colors = plt.cm.tab20(np.linspace(0, 1, len(all_categories)))
        # else:
        #     # For even more categories, use a perceptually uniform colormap
        # colors = plt.cm.hsv(np.linspace(0, 0.9, len(all_categories)))
    
        # Create stacked bars
        bottom_true = 0
        bottom_false = 0
    
        # Keep track of which categories we've already added to legend
        legend_added = set()
    
        for i, category in enumerate(all_categories):
            # Determine if this category should get a legend label
            should_label = category not in legend_added
        
            # Bar for archaeology = true
            if true_values[i] > 0:
                ax.bar(x_positions[0], true_values[i], bar_width, 
                       bottom=bottom_true, label=category if should_label else "",
                       color=colors[i], edgecolor='white', linewidth=1)
                # Add count text in the middle of the segment
                segment_center = bottom_true + true_values[i] / 2
                ax.text(x_positions[0], segment_center, str(true_values[i]), 
                       ha='center', va='center', fontweight='bold', fontsize=10 if true_values[i] > 2 else 6,
                       color='black')
                bottom_true += true_values[i]
                if should_label:
                    legend_added.add(category)

            should_label = category not in legend_added
        
            # Bar for archaeology = false
            if false_values[i] > 0:
                ax.bar(x_positions[1], false_values[i], bar_width, 
                       bottom=bottom_false, label=category if should_label else "",
                       color=colors[i], edgecolor='white', linewidth=1)
                # Add count text in the middle of the segment
                segment_center = bottom_false + false_values[i] / 2
                ax.text(x_positions[1], segment_center, str(false_values[i]), 
                       ha='center', va='center', fontweight='bold', fontsize=10 if false_values[i] > 2 else 6,
                       color='black')
                bottom_false += false_values[i]
                if should_label:
                    legend_added.add(category)
    
        # Customize the plot
        # ax.set_xlabel('Arqueología Status', fontsize=12)
        ax.set_ylabel('Número de artículos', fontsize=12)
        # ax.set_title('Distribution of Categories by Archaeology Status', fontsize=14, fontweight='bold')
        ax.set_xticks(x_positions)
        ax.set_xticklabels(bar_labels)
    
        # Add legend
        ax.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
    
        # Add value labels on bars
        if bottom_true > 0:
            ax.text(x_positions[0], bottom_true + 0.1, str(bottom_true), 
                    ha='center', va='bottom', fontweight='bold')
        if bottom_false > 0:
            ax.text(x_positions[1], bottom_false + 0.1, str(bottom_false), 
                    ha='center', va='bottom', fontweight='bold')
    
        # Adjust layout to prevent legend cutoff
        plt.tight_layout()
    
        # Print summary statistics
        print("Summary:")
        print(f"Articles with arqueología = True: {len(archaeology_true)}")
        print(f"Articles with arqueología = False/Missing: {len(archaeology_false)}")
        print(f"Categories found: {', '.join(all_categories)}")
        print("\nCategory breakdown:")
        print("Arqueología = True:")
        for cat, count in sorted(true_counts.items()):
            print(f"  {cat}: {count}")
        print("Arqueología = False/Missing:")
        for cat, count in sorted(false_counts.items()):
            print(f"  {cat}: {count}")
    
        return fig, ax

    fig, ax = create_archaeology_chart(data)

    plt.savefig('archaeology_categories.png', dpi=300)
    plt.savefig('archaeology_categories.svg', dpi=300)
    return


@app.cell
def _(data):
    import pandas as pd
    from docx import Document
    import datetime

    def format_date(date_obj):
        """Format date objects to string if needed"""
        if isinstance(date_obj, datetime.date):
            return date_obj.strftime('%d-%m-%Y')
        return str(date_obj)

    def format_list(list_item):
        """Format list items to comma-separated strings"""
        if isinstance(list_item, list):
            return ", ".join(list_item)
        return str(list_item)

    def toml_to_word_table(data, output_filename="articles_table.docx"):
        """Convert TOML data to a Word document table"""
        # Create a new Word document
        doc = Document()

        # Add a title
        doc.add_heading('Articles Table', 0)

        # Create a list to store the articles data
        articles_data = []

        # Extract data from the TOML structure
        if 'articles' in data:
            for article in data['articles']:
                # Create a dictionary for each article with all its properties
                article_dict = {}

                # Add all available fields
                for key, value in article.items():
                    if key == 'date':
                        article_dict[key] = format_date(value)
                    elif isinstance(value, list):
                        article_dict[key] = format_list(value)
                    elif isinstance(value, bool):
                        article_dict[key] = "Sí" if value else "No"
                    else:
                        article_dict[key] = str(value)

                articles_data.append(article_dict)

        # Create a pandas DataFrame from the articles data
        if articles_data:
            # Get all unique keys from all dictionaries to ensure we capture all columns
            all_keys = set()
            for article_dict in articles_data:
                all_keys.update(article_dict.keys())

            # Create a list of records with consistent keys
            complete_records = []
            for article_dict in articles_data:
                record = {key: article_dict.get(key, "") for key in all_keys}
                complete_records.append(record)

            # Create DataFrame
            df = pd.DataFrame(complete_records)

            # Add table to document
            table = doc.add_table(rows=1, cols=len(df.columns))
            table.style = 'Table Grid'

            # Add headers
            header_cells = table.rows[0].cells
            for i, column in enumerate(df.columns):
                header_cells[i].text = column.capitalize()

            # Add data rows
            for _, row in df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value) if pd.notna(value) else ""

        # Also create a separate figures table if figures exist
        figures_data = []
        if 'figures' in data:
            doc.add_heading('Figures Table', 1)

            for figure in data['figures']:
                figure_dict = {}

                for key, value in figure.items():
                    if key == 'date':
                        figure_dict[key] = format_date(value)
                    elif isinstance(value, list):
                        figure_dict[key] = format_list(value)
                    elif isinstance(value, bool):
                        figure_dict[key] = "Sí" if value else "No"
                    else:
                        figure_dict[key] = str(value)

                figures_data.append(figure_dict)

            if figures_data:
                # Get all unique keys for figures
                all_fig_keys = set()
                for figure_dict in figures_data:
                    all_fig_keys.update(figure_dict.keys())

                # Create complete records
                complete_fig_records = []
                for figure_dict in figures_data:
                    record = {key: figure_dict.get(key, "") for key in all_fig_keys}
                    complete_fig_records.append(record)

                # Create DataFrame for figures
                fig_df = pd.DataFrame(complete_fig_records)

                # Add figures table
                fig_table = doc.add_table(rows=1, cols=len(fig_df.columns))
                fig_table.style = 'Table Grid'

                # Add headers
                fig_header_cells = fig_table.rows[0].cells
                for i, column in enumerate(fig_df.columns):
                    fig_header_cells[i].text = column.capitalize()

                # Add data rows
                for _, row in fig_df.iterrows():
                    row_cells = fig_table.add_row().cells
                    for i, value in enumerate(row):
                        row_cells[i].text = str(value) if pd.notna(value) else ""

        # Save the document
        doc.save(output_filename)
        print(f"Document saved as {output_filename}")

        # If you want to also create CSV files for easy import into other tools
        if articles_data:
            pd.DataFrame(complete_records).to_csv('articles.csv', index=False)
            print("CSV file 'articles.csv' created")

        if figures_data:
            pd.DataFrame(complete_fig_records).to_csv('figures.csv', index=False)
            print("CSV file 'figures.csv' created")

    # Example usage
    # Assuming 'data' contains your parsed TOML data
    toml_to_word_table(data)

    return


if __name__ == "__main__":
    app.run()
